import warnings
warnings.filterwarnings('ignore')



import streamlit as st
import pandas as pd
import plotly.express as px
import plotly.graph_objects as go
from datetime import datetime
import urllib.parse
import re
import requests
import os

# ==============================================================================
# MODULO DE GENERACION DE INFORMES EN WORD (ANTIPASTO - ZERO EMOJIS)
# ==============================================================================
def generar_documento_word_informe(df_informe, titulo_informe="Informe de Seguimiento Proyectos SPAE"):
    import docx
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml import parse_xml
    from docx.oxml.ns import nsdecls
    import io

    doc = Document()
    for s in doc.sections:
        s.top_margin = Inches(1)
        s.bottom_margin = Inches(1)
        s.left_margin = Inches(1)
        s.right_margin = Inches(1)

    p_t = doc.add_paragraph()
    p_t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_t = p_t.add_run(titulo_informe.upper())
    r_t.font.name = 'Antipasto'
    r_t.font.size = Pt(20)
    r_t.font.bold = True
    r_t.font.color.rgb = RGBColor(0x00, 0x33, 0x66)

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_sub = p_sub.add_run("SUBDIRECCION DE PREVENCION Y ATENCION DE EMERGENCIAS - SPAE 2026")
    r_sub.font.name = 'Antipasto'
    r_sub.font.size = Pt(12)
    r_sub.font.color.rgb = RGBColor(0x00, 0x66, 0x99)

    doc.add_paragraph()

    # SECTION 1: PROJ SUMMARY
    h1 = doc.add_heading(level=1)
    r1 = h1.add_run("1. Consolidador de Estado y Avance de Proyectos")
    r1.font.name = 'Antipasto'
    r1.font.size = Pt(14)
    r1.font.bold = True
    r1.font.color.rgb = RGBColor(0x00, 0x33, 0x66)

    table = doc.add_table(rows=1, cols=6)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    hdr_cells = table.rows[0].cells
    cols = ["Codigo", "Municipio", "Linea", "Estado", "Responsable", "Avance %"]
    for i, col_name in enumerate(cols):
        cell = hdr_cells[i]
        shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="003366"/>')
        cell._tc.get_or_add_tcPr().append(shd)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(col_name)
        r.font.name = 'Antipasto'
        r.font.size = Pt(10)
        r.font.bold = True
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    for _, row in df_informe.iterrows():
        row_cells = table.add_row().cells
        vals = [
            str(row.get("Consecutivo Interno", row.get("CODIGO SIGESPLAN", ""))),
            str(row.get("MUNICIPIO", "")),
            str(row.get("LINEA INVERSION", row.get("TIPO DE PROYECTO", ""))),
            str(row.get("ESTADO", "")),
            str(row.get("Profesional SPAE", "")),
            f"{row.get('% Avance TOTAL', row.get('AVANCE', 0))}%"
        ]
        for i, val in enumerate(vals):
            cell = row_cells[i]
            p = cell.paragraphs[0]
            if i in [0, 5]:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(val)
            r.font.name = 'Antipasto'
            r.font.size = Pt(9)

    doc.add_paragraph()

    # SECTION 2: DIRECTORY TABLE
    h2 = doc.add_heading(level=1)
    r2 = h2.add_run("2. Directorio Institucional de Contactos Municipales")
    r2.font.name = 'Antipasto'
    r2.font.size = Pt(14)
    r2.font.bold = True
    r2.font.color.rgb = RGBColor(0x00, 0x33, 0x66)

    dir_table = doc.add_table(rows=1, cols=5)
    dir_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    dir_hdr = dir_table.rows[0].cells
    d_cols = ["Municipio", "Alcalde(sa)", "Secretaria / Enlace", "Correo Electronico", "Telefono"]
    for i, col_name in enumerate(d_cols):
        cell = dir_hdr[i]
        shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="003366"/>')
        cell._tc.get_or_add_tcPr().append(shd)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(col_name)
        r.font.name = 'Antipasto'
        r.font.size = Pt(10)
        r.font.bold = True
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    for _, row in df_informe.iterrows():
        row_cells = dir_table.add_row().cells
        vals = [
            str(row.get("MUNICIPIO", "")),
            str(row.get("Nombre Alcalde", "No Registrado")),
            str(row.get("Nombre Enlace Contacto", "Secretaria de Planeacion")),
            str(row.get("Correos Alcaldia", "No Registrado")),
            str(row.get("Telefono Alcaldia", "No Registrado"))
        ]
        for i, val in enumerate(vals):
            cell = row_cells[i]
            p = cell.paragraphs[0]
            if i in [0]:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(val)
            r.font.name = 'Antipasto'
            r.font.size = Pt(8.5)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


import hashlib
from auth import SUPABASE_URL, SUPABASE_KEY, save_project_to_db, log_revision_supabase, sign_in_user

st.set_page_config(page_title="Registro y Seguimiento ISC SPAE", layout="wide")

hide_st_style = """
            <style>
            #MainMenu {visibility: hidden;}
            footer {visibility: hidden;}
            header {visibility: hidden;}
            .viewerBadge_container__1JCIV {display: none !important;}
            .viewerBadge_link__1S137 {display: none !important;}
            [data-testid="stBottom"] {display: none !important;}
            .block-container {
                padding-top: 2.5rem;
                padding-bottom: 0rem;
            }
            [data-testid="stAppViewContainer"] {
                background-color: #111827 !important;
            }
            </style>
            """
st.markdown(hide_st_style, unsafe_allow_html=True)

# PROTECCION DE AUTENTICACION Y ROLES - AUTO-LOGIN ACTIVO
user = st.session_state.get("auth_user")
if not user:
    from auth import DummyUser
    user = DummyUser(email="cesar.giraldo@unidadvictimas.gov.co", role="admin")
    st.session_state["auth_user"] = user
    st.session_state["user_role"] = "admin"


user_role = st.session_state.get("user_role", "free")
user_email = getattr(user, "email", "").lower()

# 1. FUNCIONES DE SUPABASE Y LOCALES
import json

def get_local_history():
    hist_path = r"C:\Users\cagch\Desktop\senador\upv2026\junio\historial_cambios.json"
    if not os.path.exists(hist_path): return pd.DataFrame()
    try:
        with open(hist_path, "r", encoding="utf-8") as f:
            data = json.load(f)
            if data: return pd.DataFrame(data)
    except Exception: pass
    return pd.DataFrame()

def save_local_history(resumen, usuario):
    hist_path = r"C:\Users\cagch\Desktop\senador\upv2026\junio\historial_cambios.json"
    nuevo_registro = {
        "Fecha": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
        "Usuario": usuario,
        "Resumen de Cambios": resumen
    }
    data = []
    if os.path.exists(hist_path):
        try:
            with open(hist_path, "r", encoding="utf-8") as f: data = json.load(f)
        except Exception: pass
    data.insert(0, nuevo_registro)
    data = data[:50]
    try:
        with open(hist_path, "w", encoding="utf-8") as f: json.dump(data, f, indent=4, ensure_ascii=False)
    except Exception: pass

def get_all_proyectos():
    user = st.session_state.get("auth_user")
    if not user or not SUPABASE_URL or not SUPABASE_KEY: return pd.DataFrame()
    
    headers = {"apikey": SUPABASE_KEY, "Authorization": f"Bearer {getattr(user, 'access_token', SUPABASE_KEY)}"}
    
    df_result = pd.DataFrame()
    
    # 1. Intentar cargar spae_snapshot_actual
    try:
        url = f"{SUPABASE_URL}/rest/v1/proyectos?nombre_proyecto=eq.spae_snapshot_actual&select=estado_json&order=created_at.desc&limit=1"
        res = requests.get(url, headers=headers)
        if res.status_code == 200:
            data = res.json()
            if data and len(data) > 0:
                estado = data[0].get("estado_json", "[]")
                if isinstance(estado, str): estado = json.loads(estado)
                df_result = pd.DataFrame(estado)
    except Exception: pass

    # 2. Fallback a SPAE_GLOBAL_DB
    if df_result.empty:
        try:
            url = f"{SUPABASE_URL}/rest/v1/proyectos?nombre_proyecto=eq.SPAE_GLOBAL_DB&select=estado_json&order=created_at.desc&limit=1"
            res = requests.get(url, headers=headers)
            if res.status_code == 200:
                data = res.json()
                if data and len(data) > 0:
                    estado = data[0].get("estado_json", "[]")
                    if isinstance(estado, str): estado = json.loads(estado)
                    df_result = pd.DataFrame(estado)
        except Exception: pass
    
    # 3. Fallback a archivos locales
    if df_result.empty:
        try:
            import glob
            import os
            versiones_dir = r"C:\Users\cagch\Desktop\senador\upv2026\junio\Versiones_Anteriores"
            archivos = glob.glob(os.path.join(versiones_dir, "Portafolio_SPAE_Revision_*.xlsx"))
            if archivos:
                archivos.sort(key=os.path.getmtime, reverse=True)
                ultimo_archivo = archivos[0]
                df_result = leer_portafolio_nuevo(ultimo_archivo)
        except Exception: pass
        
    # Siempre cruzar con el directorio más reciente (para inyectar los contactos que hayan sido actualizados localmente)
    if not df_result.empty:
        df_dir = leer_directorio_maestro()
        if not df_dir.empty:
            df_result = cruzar_con_directorio(df_result, df_dir)
        return df_result
    
    return pd.DataFrame()

# LECTURA DIRECTORIO MAESTRO
def leer_directorio_maestro() -> pd.DataFrame:
    ruta_dir = r"C:\Users\cagch\Desktop\senador\upv2026\junio\Matriz_Seguimiento_Proyectos_Cesar.xlsx"
    if not os.path.exists(ruta_dir): return pd.DataFrame()
    try:
        df_dir = pd.read_excel(ruta_dir)
        df_dir.columns = [str(c).strip() for c in df_dir.columns]
        cols_necesarias = {}
        for c in df_dir.columns:
            cl = str(c).lower().replace('ó','o').replace('á','a').replace('é','e').replace('í','i').replace('ú','u')
            if cl == "municipio": cols_necesarias["municipio"] = c
            elif "alcalde" in cl and "tel" in cl: cols_necesarias["telefono_alcalde"] = c
            elif "contacto" in cl or "enlace" in cl: cols_necesarias["telefono_enlace"] = c
            elif "correo" in cl: cols_necesarias["correos_alcaldia"] = c
            elif cl == "codigo sigesplan" or cl == "codigo_proyecto": cols_necesarias["codigo_proyecto"] = c
        
        df_dir_clean = pd.DataFrame()
        for k, v in cols_necesarias.items(): df_dir_clean[k] = df_dir[v]
        return df_dir_clean
    except Exception as e: return pd.DataFrame()

# 2. LECTURA Y NORMALIZACION DEL PORTAFOLIO SPAE 2026
def leer_portafolio_nuevo(archivo) -> pd.DataFrame:
    df = pd.read_excel(archivo, sheet_name="Registro Avance x Responsables", header=3, engine="openpyxl")
    # Normalizar cabeceras a minúsculas sin tildes para facilitar el mapeo
    df.columns = [str(c).lower().replace('ó','o').replace('á','a').replace('é','e').replace('í','i').replace('ú','u').replace('°','').strip() for c in df.columns]
    
    m_col = None
    for c in df.columns:
        if c == "municipio":
            m_col = c
            break
            
    if m_col:
        df = df.dropna(how="all", subset=[m_col])
    
    rename_map = {
        "consecutivo interno": "codigo_interno",
        "codigo sigesplan": "codigo_sigesplan",
        "departamento": "departamento",
        "municipio": "municipio",
        "codigo dane": "codigo_dane",
        "linea inversion": "linea_inversion",
        "tipo de proyecto": "tipo_proyecto",
        "titulo del proyecto": "titulo_proyecto",
        "estado": "estado_formulacion",
        "profesional spae": "profesional_spae",
        "profesional ggp": "profesional_ggp",
        "presupuesto proyectado": "presupuesto_proyectado",
        "presupuesto solicitado": "presupuesto_solicitado",
        "diferencia presupuesto": "diferencia_presupuesto",
        "n personas beneficiadas - victimas registradas": "beneficiados_victimas",
        "n personas beneficiadas - poblacion vulnerable": "beneficiados_vulnerables",
        "total n personas beneficiadas": "total_beneficiados"
    }
    
    for k, v in rename_map.items():
        if k in df.columns:
            df.rename(columns={k: v}, inplace=True)

    for v in rename_map.values():
        if v not in df.columns: df[v] = None

    for num_col in ["presupuesto_proyectado", "presupuesto_solicitado", "diferencia_presupuesto", 
                    "total_beneficiados", "beneficiados_victimas", "beneficiados_vulnerables"]:
        df[num_col] = pd.to_numeric(df[num_col], errors="coerce").fillna(0)

    ANEXOS_ISC_COLS = [
        "carta de intencion",
        "carta de presentacion proyecto",
        "ficha tecnica perfil de proyecto",
        "esquema financiero",
        "registro fotografico",
        "acta de socializacion y concertacion",
        "certificado de compromisos entidad",
        "certificado de compromisos comunidad",
        "certificacion usos de suelo",
        "certificacion servicios publicos",
        "certificacion vias de acceso",
        "certificado de tradicion y libertad",
        "presupuesto detallado",
        "analisis precios unitarios apu",
        "estudios y disenos"
    ]
    
    # Renombrar estáticamente los índices para F-1, F-2 y F-3
    if len(df.columns) > 54:
        map_f1 = {df.columns[i]: ANEXOS_ISC_COLS[i-40] for i in range(40, 55)}
        df.rename(columns=map_f1, inplace=True)
    if len(df.columns) > 67:
        map_f2 = {df.columns[i]: f"anexo_agro_{i-55}" for i in range(55, 68)}
        df.rename(columns=map_f2, inplace=True)
    if len(df.columns) > 73:
        map_f3 = {df.columns[i]: f"anexo_dmec_{i-68}" for i in range(68, 74)}
        df.rename(columns=map_f3, inplace=True)
    
    ESTADOS_VALIDOS = {"recibido", "no recibido", "en ajustes", "en proceso"}

    def calcular_pct_anexos_isc(row: pd.Series) -> float:
        tipo = str(row.get('tipo_proyecto', '')).lower() + " " + str(row.get('linea_inversion', '')).lower()
        
        # Determinar el rango de columnas a evaluar según el tipo de proyecto
        if 'agro' in tipo or 'productivo' in tipo:
            cols_idx = range(55, 68)  # F-2 AGRO
        elif 'dotacion' in tipo or 'mobiliari' in tipo or 'dmec' in tipo:
            cols_idx = range(68, 74)  # F-3 DMEC
        else: 
            cols_idx = range(40, 55)  # F-1 ISC (por defecto)
            
        estados = []
        for idx in cols_idx:
            if idx < len(df.columns):
                col_name = df.columns[idx]
                if col_name not in row.index:
                    continue
                val = str(row[col_name]).strip().lower()
                if not val or val == "no aplica" or val == "nan":
                    continue
                if val in ESTADOS_VALIDOS:
                    estados.append(val)
        
        if not estados:
            return 0.0
            
        total_aplicables = len(estados)
        total_recibidos = sum(1 for v in estados if v == "recibido")
        return round((total_recibidos / total_aplicables) * 100.0, 1)

    df["pct_anexos_isc"] = df.apply(calcular_pct_anexos_isc, axis=1)
    
    def clasificar_semaforo_isc(pct: float) -> str:
        if pct >= 80: return "Alto (>=80%)"
        elif pct >= 40: return "Medio (40-79%)"
        else: return "Bajo (<40%)"
        
    df["categoria_isc"] = df["pct_anexos_isc"].apply(clasificar_semaforo_isc)
    
    df["codigo_proyecto"] = df.apply(
        lambda r: str(r.get("codigo_sigesplan")).strip() if pd.notna(r.get("codigo_sigesplan")) and str(r.get("codigo_sigesplan")).strip() != "" 
        else str(r.get("codigo_interno")).strip() if pd.notna(r.get("codigo_interno")) and str(r.get("codigo_interno")).strip() != ""
        else f"PRJ-{str(r.get('municipio', 'XXX')).upper()[:5]}-{abs(hash(str(r.get('municipio', '')))) % 10000}", axis=1
    )
    
    # Compatibilidad para campos obligatorios en otras secciones
    df["nombre_proyecto"] = df["titulo_proyecto"]
    df["estado"] = df["estado_formulacion"]
    
    return df

# CRUCE JERARQUICO CON DIRECTORIO
def cruzar_con_directorio(df_port: pd.DataFrame, df_dir: pd.DataFrame) -> pd.DataFrame:
    if df_dir.empty:
        df_port["match_type"] = "Sin Directorio"
        return df_port
    
    df_port["muni_key"] = df_port["municipio"].astype(str).str.lower().str.strip()
    col_cod_dir = "codigo_proyecto" if "codigo_proyecto" in df_dir.columns else "codigo_proyecto_dir"
    col_muni_dir = "municipio" if "municipio" in df_dir.columns else "municipio_dir"
    col_tel_a_dir = "telefono_alcalde" if "telefono_alcalde" in df_dir.columns else "telefono_alcalde_dir"
    col_tel_e_dir = "telefono_enlace" if "telefono_enlace" in df_dir.columns else "telefono_enlace_dir"
    col_cor_dir = "correos_alcaldia" if "correos_alcaldia" in df_dir.columns else "correos_alcaldia_dir"

    df_dir_copy = df_dir.copy()
    if col_muni_dir in df_dir_copy.columns:
        df_dir_copy["muni_key"] = df_dir_copy[col_muni_dir].astype(str).str.lower().str.strip()
    
    match_status, tel_alcalde, tel_enlace, correo_alcaldia = [], [], [], []
    
    for _, row in df_port.iterrows():
        cod, mkey = str(row.get("codigo_proyecto", "")).strip(), row.get("muni_key", "")
        match_row, m_type = None, "Sin Match"
        
        if col_cod_dir in df_dir_copy.columns and cod and cod != "nan":
            filtro = df_dir_copy[df_dir_copy[col_cod_dir].astype(str).str.strip() == cod]
            if not filtro.empty: match_row, m_type = filtro.iloc[0], "Exacto (Codigo)"
        
        if match_row is None and "muni_key" in df_dir_copy.columns and mkey and mkey != "nan":
            filtro = df_dir_copy[df_dir_copy["muni_key"] == mkey]
            if not filtro.empty: match_row, m_type = filtro.iloc[0], "Aproximado (Municipio)"
                
        if match_row is not None:
            match_status.append(m_type)
            tel_alcalde.append(row.get("telefono_alcalde") if pd.notna(row.get("telefono_alcalde")) else match_row.get(col_tel_a_dir))
            tel_enlace.append(row.get("telefono_enlace") if pd.notna(row.get("telefono_enlace")) else match_row.get(col_tel_e_dir))
            correo_alcaldia.append(row.get("correos_alcaldia") if pd.notna(row.get("correos_alcaldia")) else match_row.get(col_cor_dir))
        else:
            match_status.append(m_type)
            tel_alcalde.append(row.get("telefono_alcalde"))
            tel_enlace.append(row.get("telefono_enlace"))
            correo_alcaldia.append(row.get("correos_alcaldia"))
            
    df_port["match_type"] = match_status
    df_port["telefono_alcalde"] = tel_alcalde
    df_port["telefono_enlace"] = tel_enlace
    df_port["correos_alcaldia"] = correo_alcaldia
    return df_port


def comparar_snapshot_vs_nuevo(up_dict, db_dict):
    nuevos = 0
    actualizados = 0
    sin_cambios = 0
    
    docs_subidos = []
    pct_avance_cambios = []
    obs_nuevas = []
    detalle_cambios = []
    
    ANEXOS_DOCS = [
        "carta de intencion", "carta de presentacion proyecto", "ficha tecnica perfil de proyecto",
        "esquema financiero", "registro fotografico", "acta de socializacion y concertacion",
        "certificado de compromisos entidad", "certificado de compromisos comunidad", "certificacion usos de suelo",
        "certificacion servicios publicos", "certificacion vias de acceso", "certificado de tradicion y libertad",
        "presupuesto detallado", "analisis precios unitarios apu", "estudios y disenos"
    ]
    ESTADOS_MEJORA = ["recibido", "en proceso", "ok", "si", "1"]
    
    for k, v_up in up_dict.items():
        municipio_nombre = str(v_up.get("municipio", k))
        profesional_spae = str(v_up.get("profesional_spae", "Sin SPAE")).strip()
        if profesional_spae == "nan": profesional_spae = "Sin SPAE"
        profesional_ggp = str(v_up.get("profesional_ggp", "Sin GGP")).strip()
        if profesional_ggp == "nan": profesional_ggp = "Sin GGP"
        
        profesional = f"SPAE: {profesional_spae} | GGP: {profesional_ggp}"
        
        if k in db_dict:
            changed = False
            cambios_proyecto = {}
            for col_name, val in v_up.items():
                if val is not None:
                    val_old = db_dict[k].get(col_name)
                    if str(val_old) != str(val) and col_name not in ["muni_key", "match_type", "telefono_alcalde", "telefono_enlace", "correos_alcaldia"]:
                        changed = True
                        cambios_proyecto[col_name] = {"old": val_old, "new": val}
                        
                        if col_name in ANEXOS_DOCS or str(col_name).startswith("anexo_agro_") or str(col_name).startswith("anexo_dmec_"):
                            v_str = str(val).lower().strip()
                            o_str = str(val_old).lower().strip()
                            if v_str in ESTADOS_MEJORA and o_str not in ESTADOS_MEJORA:
                                docs_subidos.append({"muni": municipio_nombre, "prof": profesional, "doc": col_name})
                                
                        if col_name == "pct_anexos_isc":
                            try:
                                v_f = float(val)
                                o_f = float(val_old or 0)
                                if v_f != o_f:
                                    pct_avance_cambios.append({"muni": municipio_nombre, "prof": profesional, "old": o_f, "new": v_f})
                            except: pass
                            
                        if "observacion" in str(col_name).lower():
                            v_str = str(val).strip()
                            o_str = str(val_old).strip()
                            if v_str and v_str not in ["nan", "None", ""] and v_str != o_str:
                                obs_nuevas.append({"muni": municipio_nombre, "prof": profesional, "col": col_name, "obs": v_str})
                                
                    db_dict[k][col_name] = val
            
            if changed:
                actualizados += 1
                detalle_cambios.append({"proyecto": municipio_nombre, "codigo": k, "cambios": cambios_proyecto})
            else:
                sin_cambios += 1
        else:
            db_dict[k] = v_up
            nuevos += 1
            detalle_cambios.append({"proyecto": municipio_nombre, "codigo": k, "cambios": "NUEVO PROYECTO"})
            
    resumen_funcional = []
    if nuevos > 0: resumen_funcional.append(f"Se identificaron y agregaron {nuevos} proyectos nuevos.")
    resumen_funcional.append(f"Se actualizaron {actualizados} proyectos existentes.")
    
    if docs_subidos:
        resumen_funcional.append("**Documentos subidos a estado Recibido/Proceso:**")
        for d in docs_subidos: resumen_funcional.append(f"  - **{d['muni']}** ({d['prof']}): se subió *{d['doc']}*")
            
    if pct_avance_cambios:
        resumen_funcional.append("**Cambios en % de Avance de Anexos:**")
        for p in pct_avance_cambios:
            signo = "subió" if p['new'] > p['old'] else "bajó"
            resumen_funcional.append(f"  - **{p['muni']}** ({p['prof']}): {signo} de {p['old']}% a {p['new']}%")
            
    if obs_nuevas:
        resumen_funcional.append("**Nuevas Observaciones Detectadas:**")
        for o in obs_nuevas: resumen_funcional.append(f"  - **{o['muni']}** ({o['prof']}) en *{o['col']}*: \"{o['obs']}\"")
            
    return db_dict, detalle_cambios, resumen_funcional, (nuevos + actualizados)


# 3. UPSERT Y DETECCION DE CAMBIOS
def upsert_proyectos(df_upload_parsed: pd.DataFrame, df_db: pd.DataFrame, usuario: str, uploaded_file):
    user = st.session_state.get("auth_user")
    if not user: return 0, []
    
    file_bytes = uploaded_file.getvalue()
    hash_archivo = hashlib.md5(file_bytes).hexdigest()
    
    resumen_funcional = []
    detalle_cambios = []
    filas = 0
    
    total_proyectos = len(df_upload_parsed)
    presupuesto_total = float(df_upload_parsed["presupuesto_proyectado"].sum()) if "presupuesto_proyectado" in df_upload_parsed else 0.0
    beneficiados_totales = float(df_upload_parsed["total_beneficiados"].sum()) if "total_beneficiados" in df_upload_parsed else 0.0
    
    if not df_db.empty and "codigo_proyecto" in df_db.columns and "codigo_proyecto" in df_upload_parsed.columns:
        for col in df_upload_parsed.columns:
            df_upload_parsed[col] = df_upload_parsed[col].apply(
                lambda x: None if pd.isna(x) or str(x).lower() in ['nan', 'inf', '-inf'] else (x.strftime('%Y-%m-%d %H:%M:%S') if hasattr(x, 'strftime') else x)
            )
        db_dict = {str(r.get("codigo_proyecto")): r for r in df_db.to_dict("records")}
        up_dict = {str(r.get("codigo_proyecto")): r for r in df_upload_parsed.to_dict("records")}
        
        db_dict, detalle_cambios, resumen_funcional, filas = comparar_snapshot_vs_nuevo(up_dict, db_dict)
        registros = list(db_dict.values())
        
    else:
        df_upload = df_upload_parsed.copy()
        for col in df_upload.columns:
            df_upload[col] = df_upload[col].apply(
                lambda x: None if pd.isna(x) or str(x).lower() in ['nan', 'inf', '-inf'] else (x.strftime('%Y-%m-%d %H:%M:%S') if hasattr(x, 'strftime') else x)
            )
        registros_raw = df_upload.to_dict("records")
        registros = [{k: v for k, v in d.items() if v is not None} for d in registros_raw]
        resumen_funcional.append(f"Carga inicial de {len(registros)} proyectos.")
        nuevos = len(registros)
        actualizados = 0
        sin_cambios = 0
        detalle_cambios = [{"msg": "Carga inicial masiva"}]
    
    try:
        # Guardar snapshot principal
        save_project_to_db(user, "spae_snapshot_actual", usuario, "Directorio SPAE", "N/A", registros)
        
        # Guardar log en Supabase (si la tabla existe)
        try:
            revision_num = int(datetime.now().timestamp())
            log_revision_supabase(
                user=user, 
                archivo=uploaded_file.name, 
                revisado_por=usuario, 
                hash_archivo=hash_archivo, 
                total_proyectos=total_proyectos, 
                presupuesto_total=presupuesto_total, 
                beneficiados_totales=beneficiados_totales, 
                detalle_cambios={"resumen": resumen_funcional, "detalle": detalle_cambios}, 
                revision_num=revision_num
            )
        except Exception as e_log:
            st.warning(f"Nota: El snapshot se actualizó, pero el log de revisión falló. ¿Aseguraste crear la tabla `spae_portafolio_revisiones`? Error: {e_log}")
        
        resumen_texto = "\n".join(resumen_funcional)
        save_local_history(resumen_texto, usuario)
        
        return len(registros), resumen_funcional
    except Exception as e:
        st.error(f"Error al guardar snapshot con Supabase: {str(e)}")
        return 0, []

def obtener_faltantes_texto(row) -> str:
    tipo = str(row.get('tipo_proyecto', '')).lower() + " " + str(row.get('linea_inversion', '')).lower()
    
    if 'agro' in tipo or 'productivo' in tipo:
        posibles = [f"anexo_agro_{i-55}" for i in range(55, 68)]
    elif 'dotacion' in tipo or 'mobiliari' in tipo or 'dmec' in tipo:
        posibles = [f"anexo_dmec_{i-68}" for i in range(68, 74)]
    else: 
        posibles = ["carta de intencion", "carta de presentacion proyecto", "ficha tecnica perfil de proyecto", "esquema financiero", "registro fotografico", "acta de socializacion y concertacion", "certificado de compromisos entidad", "certificado de compromisos comunidad", "certificacion usos de suelo", "certificacion servicios publicos", "certificacion vias de acceso", "certificado de tradicion y libertad", "presupuesto detallado", "analisis precios unitarios apu", "estudios y disenos"]
        
    faltantes = []
    for col in posibles:
        if col in row.index:
            val = str(row[col]).strip().lower()
            if val in {"no recibido", "pendiente", "en ajustes"}:
                faltantes.append(col.replace("_", " ").title())
                
    if not faltantes:
        return "Según nuestra matriz, no hay anexos pendientes registrados."
        
    listado = "\n".join(f"- {anexo}" for anexo in faltantes)
    
    texto = (
        "Nos encontramos a la espera de la siguiente documentación faltante "
        "para avanzar con su proyecto SPAE:\n\n"
        f"{listado}\n\n"
    )
        
    texto += (
        "Si ya envió alguno de estos documentos, por favor reenvíelos a mi correo "
        "personal para consolidar la información: civcesar2021@gmail.com.\n"
        "Es importante que queden asociados a este proyecto y municipio."
    )
    
    return texto

def whatsapp_link_cascada(row) -> tuple:
    telefono = row.get("telefono_enlace") or row.get("telefono_alcalde")
    if not pd.notna(telefono) or str(telefono).strip() in ["", "0", "None"]:
        return None, "Sin Contacto"

    municipio = row.get("municipio", "")
    proyecto = row.get("titulo_proyecto", "")
    faltantes = obtener_faltantes_texto(row)

    mensaje = (
        f"Cordial saludo,\n\n"
        f"Le escribo respecto al proyecto SPAE \"{proyecto}\" en el municipio de {municipio}.\n\n"
        f"{faltantes}\n\n"
        "Ing. Msc. Cesar Giraldo — Profesional SPAE."
    )

    tel_limpio = "".join(ch for ch in str(telefono) if ch.isdigit())
    if len(tel_limpio) >= 10:
        tel = tel_limpio[:10]
        if len(tel) == 10: tel = "57" + tel
        tipo = "Enlace" if pd.notna(row.get("telefono_enlace")) and str(row.get("telefono_enlace")).strip() not in ["", "0", "None"] else "Alcalde"
        return f"https://wa.me/{tel}?text={urllib.parse.quote(mensaje)}", tipo

    return None, "Sin Contacto"

def email_link(row) -> str:
    correo = row.get("correos_alcaldia")
    if pd.isna(correo) or str(correo).strip() in ["", "None"]: return None
    email_limpio = str(correo).split()[0].split(',')[0].strip()
    if "@" not in email_limpio: return None
    
    municipio = row.get("municipio", "")
    proyecto = row.get("titulo_proyecto", "")
    faltantes = obtener_faltantes_texto(row)

    asunto = f"Documentación pendiente proyecto SPAE — {municipio}"
    cuerpo = (
        f"Cordial saludo,\n\n"
        f"Le escribo respecto al proyecto SPAE \"{proyecto}\" en el municipio de {municipio}.\n\n"
        f"{faltantes}\n\n"
        "Quedo atento.\n\n"
        "Ing. Msc. Cesar Giraldo\n"
        "Profesional SPAE"
    )

    return f"https://mail.google.com/mail/?view=cm&fs=1&to={email_limpio}&su={urllib.parse.quote(asunto)}&body={urllib.parse.quote(cuerpo)}"


# ==========================================
# INTERFAZ DE USUARIO
# ==========================================

st.title("Seguimiento Estratégico de Portafolio SPAE 2026")

versiones_dir = r"C:\Users\cagch\Desktop\senador\upv2026\junio\Versiones_Anteriores"
resumen_file = os.path.join(versiones_dir, "Ultimo_Resumen.txt")
if os.path.exists(resumen_file):
    with st.expander("Ver Último Reporte de Cambios (Histórico)", expanded=False):
        try:
            with open(resumen_file, "r", encoding="utf-8") as f:
                ultimo_res = f.read()
            st.markdown(ultimo_res)
            st.download_button("Descargar Reporte (.txt)", data=ultimo_res.encode('utf-8'), file_name="Ultimo_Reporte_Cambios.txt", mime="text/plain")
        except: pass


df_db = get_all_proyectos()

if user_role == "admin":
    with st.expander("Subir Nuevo Portafolio / Actualizar Matriz"):
        st.info("Carga el archivo Excel oficial. El sistema calculará diferencias y guardará la revisión.")
        uploaded_file = st.file_uploader("Arrastra aquí tu matriz (.xlsx)", type=["xlsx"])
        
        if uploaded_file:
            df_upload_parsed = leer_portafolio_nuevo(uploaded_file)
            
            # CRUZAMOS CON EL DIRECTORIO LOCAL PARA TRAER TELEFONOS
            df_dir = leer_directorio_maestro()
            df_upload_parsed = cruzar_con_directorio(df_upload_parsed, df_dir)
            
            col_diag1, col_diag2 = st.columns(2)
            with col_diag1:
                st.markdown("**Diagnóstico de Ingesta**")
                st.write(f"Proyectos detectados: {len(df_upload_parsed)}")
            with col_diag2:
                st.markdown("**Cruce con Directorio CRM**")
                if "match_type" in df_upload_parsed.columns:
                    st.dataframe(df_upload_parsed["match_type"].value_counts().reset_index(), hide_index=True)
                
            st.warning("Para aplicar los cambios y verlos en los tableros, haz clic en el botón de abajo:")
            if st.button("Guardar Versión y Analizar Cambios", type="primary", use_container_width=True):
                with st.spinner("Comparando con el snapshot anterior y guardando revisión..."):
                    usuario_email = getattr(st.session_state.get("auth_user"), "email", "admin")
                    filas, resumen = upsert_proyectos(df_upload_parsed, df_db, usuario_email, uploaded_file)
                    
                    # Guardado físico local
                    versiones_dir = r"C:\Users\cagch\Desktop\senador\upv2026\junio\Versiones_Anteriores"
                    if not os.path.exists(versiones_dir): os.makedirs(versiones_dir)
                    timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
                    file_path = os.path.join(versiones_dir, f"Portafolio_SPAE_Revision_{timestamp}.xlsx")
                    try:
                        with open(file_path, "wb") as f: f.write(uploaded_file.getbuffer())
                    except: pass
                    
                    if filas > 0:
                        st.success("¡Base de datos actualizada exitosamente!")
                        if resumen:
                            # Guardar resumen en txt para mostrarlo en el futuro
                            try:
                                resumen_texto = "### Último Resumen de Cambios Funcionales:\n\n" + "\n".join([f"- {r}" for r in resumen])
                                with open(os.path.join(versiones_dir, "Ultimo_Resumen.txt"), "w", encoding="utf-8") as f:
                                    f.write(resumen_texto)
                            except: pass
                            
                            st.markdown("### Resumen de Cambios Funcionales detectados:")
                            for r in resumen:
                                st.markdown(f"- {r}")
                    else:
                        st.warning("No se pudo guardar la versión en la nube (posible fallo de conexión o autenticación RLS), pero puedes ver los datos en los tableros de abajo en modo memoria temporal.")
            
            # SOBREESCRIBIR df_db CON LOS DATOS RECIÉN SUBIDOS
            # Esto asegura que el dashboard siempre muestre la matriz que el usuario acaba de subir,
            # independientemente de si la conexión a Supabase funcionó o falló.
            df_db = df_upload_parsed.copy()

if df_db.empty:
    df_db = leer_directorio_maestro()
    if not df_db.empty:
        st.success("La base de datos en la nube está vacía, pero he precargado tus contactos locales del Directorio Maestro.")
    else:
        st.warning("Tu base de datos está vacía. Por favor sube un archivo Excel en el panel de arriba para iniciar el ecosistema.")
        st.stop()

df = df_db.copy()

st.divider()

# SIDEBAR: FILTROS ROBUSTOS Y SEGURIDAD
st.sidebar.header("Filtros de Base de Datos")

# LÓGICA DE SEGURIDAD POR CORREO (RLS a nivel aplicación)
# Si no es admin, intentamos detectar qué municipios le pertenecen
df_dir_filtro = leer_directorio_maestro()
mis_muni = []
es_alcalde_o_enlace = False

if user_role != "admin" and user_email and not df_dir_filtro.empty:
    if "correos_alcaldia" in df_dir_filtro.columns:
        # Buscar si el correo pertenece a una alcaldía
        mask_alcalde = df_dir_filtro["correos_alcaldia"].astype(str).str.lower().apply(lambda x: user_email in x)
        mis_muni = df_dir_filtro[mask_alcalde]["municipio"].astype(str).str.lower().str.strip().tolist()
        if mis_muni: es_alcalde_o_enlace = True

# Forzar filtro si es alcalde, o dar la opción si es profesional/otro
st.sidebar.markdown("### Vistas Rápidas")
if es_alcalde_o_enlace:
    st.sidebar.success("Modo Alcaldía/Enlace Activo")
    st.sidebar.info(f"Viendo únicamente datos de: {', '.join([m.title() for m in mis_muni])}")
    df = df[df["municipio"].astype(str).str.lower().str.strip().isin(mis_muni)]
else:
    ver_mis_municipios = st.sidebar.checkbox("Filtrar Mis Municipios Asignados", value=(user_role != "admin"))
    if ver_mis_municipios and not df_dir_filtro.empty and "municipio" in df_dir_filtro.columns:
        # Comportamiento heredado para Cesar u otros profesionales locales
        mis_muni_cesar = df_dir_filtro["municipio"].astype(str).str.lower().str.strip().tolist()
        df = df[df["municipio"].astype(str).str.lower().str.strip().isin(mis_muni_cesar)]

busqueda = st.sidebar.text_input("Buscador Libre (Municipio, Proyecto, etc.)")
if busqueda:
    mask = df.astype(str).apply(lambda x: x.str.contains(busqueda, case=False, na=False)).any(axis=1)
    df = df[mask]

st.sidebar.markdown("### Filtros por Profesionales")

profesionales_spae = []
if "profesional_spae" in df.columns:
    profesionales_spae = sorted([str(x) for x in df["profesional_spae"].dropna().unique() if str(x).strip() not in ['', 'None']])

profesionales_ggp = []
if "profesional_ggp" in df.columns:
    profesionales_ggp = sorted([str(x) for x in df["profesional_ggp"].dropna().unique() if str(x).strip() not in ['', 'None']])

st.sidebar.metric("Profesionales SPAE detectados", len(profesionales_spae))
st.sidebar.metric("Profesionales GGP detectados", len(profesionales_ggp))

seleccion_spae = st.sidebar.multiselect("Profesionales SPAE", options=profesionales_spae, default=st.session_state.get("seleccion_spae", []))
st.session_state["seleccion_spae"] = seleccion_spae
seleccion_ggp = st.sidebar.multiselect("Profesionales GGP", options=profesionales_ggp, default=st.session_state.get("seleccion_ggp", []))
st.session_state["seleccion_ggp"] = seleccion_ggp

# Filter
df_filtrado = df.copy()
if seleccion_spae:
    df_filtrado = df_filtrado[df_filtrado["profesional_spae"].isin(seleccion_spae)]
if seleccion_ggp:
    df_filtrado = df_filtrado[df_filtrado["profesional_ggp"].isin(seleccion_ggp)]

df = df_filtrado



# TABS PRINCIPALES DE DASHBOARD
tab_municipio, tab_grupo, tab_crm, tab_historial = st.tabs(["Dashboard por Municipio", "Dashboard por Grupo", "Directorio CRM", "Historial de Revisiones"])

with tab_municipio:
    st.markdown("## Análisis de Portafolio por Municipio")
    
    muni_opciones = sorted([str(m) for m in df["municipio"].dropna().unique() if str(m).strip()]) if "municipio" in df.columns else []
    
    if "municipio_seleccionado" not in st.session_state:
        st.session_state["municipio_seleccionado"] = "-- Todos --"
        
    opciones_todas = ["-- Todos --"] + muni_opciones
    idx = opciones_todas.index(st.session_state["municipio_seleccionado"]) if st.session_state["municipio_seleccionado"] in opciones_todas else 0
    muni_seleccionado = st.selectbox("Seleccione un Municipio para ver el análisis detallado:", options=opciones_todas, index=idx)
    st.session_state["municipio_seleccionado"] = muni_seleccionado
    
    df_muni = df.copy()
    if muni_seleccionado != "-- Todos --":
        df_muni = df[df["municipio"] == muni_seleccionado]
        
    for col in ["presupuesto_proyectado", "presupuesto_solicitado", "total_beneficiados"]:
        if col in df_muni.columns:
            df_muni[col] = pd.to_numeric(df_muni[col], errors="coerce").fillna(0)
            
    kpi_presup_proy = df_muni.get("presupuesto_proyectado", pd.Series([0])).sum()
    kpi_presup_soli = df_muni.get("presupuesto_solicitado", pd.Series([0])).sum()
    kpi_benef_tot   = df_muni.get("total_beneficiados", pd.Series([0])).sum()
    kpi_n_proy      = len(df_muni)
    
    kpi_pct_anexos = pd.to_numeric(df_muni.get("%_anexos_isc_recibidos", pd.Series([0])), errors="coerce").mean()
    if pd.isna(kpi_pct_anexos): kpi_pct_anexos = 0.0
    
    # Panel de debug temporal
    with st.expander("Debug municipal (temporal)", expanded=True):
        st.write({
            "Municipio": muni_seleccionado,
            "Filas en df_muni": len(df_muni),
            "Suma presupuesto proyectado": float(df_muni.get("presupuesto_proyectado", pd.Series([0])).sum()),
            "Suma presupuesto solicitado": float(df_muni.get("presupuesto_solicitado", pd.Series([0])).sum()),
            "Suma total beneficiados": float(df_muni.get("total_beneficiados", pd.Series([0])).sum()),
            "Promedio pct_anexos_isc": float(df_muni.get("pct_anexos_isc", pd.Series([0])).mean()),
        })
        
    col1, col2, col3, col4 = st.columns(4)

    col1.metric("Presupuesto Proyectado", f"${df_muni.get('presupuesto_proyectado', pd.Series([0])).sum():,.0f}")
    col2.metric("Presupuesto Solicitado", f"${df_muni.get('presupuesto_solicitado', pd.Series([0])).sum():,.0f}")
    col3.metric("Beneficiados Totales", f"{int(df_muni.get('total_beneficiados', pd.Series([0])).sum()):,}")
    col4.metric("Avance Promedio (General)", f"{df_muni.get('pct_anexos_isc', pd.Series([0])).mean():.1f} %")
    
    st.divider()

    # GENERADOR DE MENSAJES AUTOMATICOS DE WHATSAPP (COBRO DE ANEXOS)
    if muni_seleccionado != "-- Todos --":
        with st.expander(" Generador Automático de Recordatorios (WhatsApp)"):
            st.info("Genera mensajes listos para enviar al Enlace o Alcalde con los documentos exactos que le faltan a cada proyecto.")
            
            # Buscar teléfono en directorio local
            df_dir_local = leer_directorio_maestro()
            telefono_muni = None
            if not df_dir_local.empty and "municipio" in df_dir_local.columns:
                df_dir_match = df_dir_local[df_dir_local["municipio"].astype(str).str.lower().str.strip() == str(muni_seleccionado).lower().strip()]
                if not df_dir_match.empty:
                    row_dir = df_dir_match.iloc[0]
                    tel = row_dir.get("telefono_enlace") if pd.notna(row_dir.get("telefono_enlace")) else row_dir.get("telefono_alcalde")
                    if pd.notna(tel) and str(tel).strip() not in ["", "0", "None", "nan"]:
                        telefono_muni = str(tel).strip()

            if telefono_muni:
                st.success(f"Contacto principal encontrado: {telefono_muni}")
                import urllib.parse
                hay_pendientes = False
                for _, row in df_muni.iterrows():
                    faltantes = []
                    for col in df_muni.columns:
                        if "anexo" in col.lower() or "cert" in col.lower() or "doc_" in col.lower():
                            if col == "pct_anexos_isc": continue
                            val = str(row[col]).strip().upper()
                            if val in ["NO", "", "NAN", "NONE", "PENDIENTE"]:
                                nombre_doc = col.replace("_", " ").title()
                                faltantes.append(nombre_doc)
                    
                    if faltantes:
                        hay_pendientes = True
                        codigo = row.get("codigo_interno", "S/N")
                        titulo = row.get("titulo_proyecto", "Sin Título")
                        mensaje = f"Cordial saludo. Desde la supervisión SPAE revisamos el proyecto *{codigo}* en el municipio de {muni_seleccionado.title()} y notamos que faltan los siguientes documentos:\n\n- " + "\n- ".join(faltantes) + "\n\nQuedamos atentos para avanzar con la viabilidad técnica."
                        url_wa = f"https://wa.me/57{telefono_muni.replace(' ', '')}?text={urllib.parse.quote(mensaje)}"
                        st.markdown(f"**Proyecto {codigo}:** {len(faltantes)} documentos faltantes. [ Enviar Recordatorio por WhatsApp]({url_wa})")
                
                if not hay_pendientes:
                    st.success("Todos los proyectos de este municipio tienen sus anexos completos.")
            else:
                st.warning("No se encontró un teléfono registrado (Alcalde o Enlace) para este municipio en el Directorio.")

    st.divider()
    
    if not df_muni.empty:
        c1, c2 = st.columns(2)
        if "titulo_proyecto" in df_muni.columns and "presupuesto_proyectado" in df_muni.columns:
            fig1 = px.bar(df_muni, x="codigo_interno", y="presupuesto_proyectado", title="Presupuesto Proyectado por Proyecto", color=df_muni.get("estado_formulacion"), hover_data=["titulo_proyecto"])
            fig1.update_layout(xaxis_title="", yaxis_title="Presupuesto")
            c1.plotly_chart(fig1, use_container_width=True)
        
        if "beneficiados_victimas" in df_muni.columns and "beneficiados_vulnerables" in df_muni.columns:
            df_pop = df_muni[["codigo_interno", "titulo_proyecto", "beneficiados_victimas", "beneficiados_vulnerables"]].melt(id_vars=["codigo_interno", "titulo_proyecto"], var_name="Población", value_name="Cantidad")
            fig2 = px.bar(df_pop, x="codigo_interno", y="Cantidad", color="Población", title="Víctimas vs Población Vulnerable por Proyecto", barmode="stack", hover_data=["titulo_proyecto"])
            fig2.update_layout(xaxis_title="")
            c2.plotly_chart(fig2, use_container_width=True)
            
        st.divider()
        
        if "pct_anexos_isc" in df_muni.columns and len(df_muni) > 0:
            import textwrap
            df_barras = df_muni[["codigo_interno", "titulo_proyecto", "pct_anexos_isc", "tipo_proyecto"]].copy()
            if "profesional_spae" in df_muni.columns:
                df_barras["prof_spae"] = df_muni["profesional_spae"]
            if "profesional_ggp" in df_muni.columns:
                df_barras["prof_ggp"] = df_muni["profesional_ggp"]
                
            df_barras["codigo"] = df_barras["codigo_interno"].astype(str)
            df_barras["nombre_completo"] = df_barras["titulo_proyecto"].apply(lambda x: '<br>'.join(textwrap.wrap(str(x), width=60)) if pd.notna(x) else "Sin Título")
        
            hover_cols = ["nombre_completo", "tipo_proyecto"]
            if "prof_spae" in df_barras.columns: hover_cols.append("prof_spae")
            if "prof_ggp" in df_barras.columns: hover_cols.append("prof_ggp")

            fig = px.bar(
                df_barras,
                x="codigo",
                y="pct_anexos_isc",
                title=f"Avance documental ISC/DMEC/AGRO por proyecto - {muni_seleccionado}",
                labels={"codigo": "Proyecto", "pct_anexos_isc": "% anexos recibidos", "nombre_completo": "Proyecto"},
                color="pct_anexos_isc",
                color_continuous_scale=["#d32f2f", "#ffb300", "#4caf50"],
                range_y=[0, 100],
                hover_data=hover_cols
            )
            fig.update_layout(xaxis_tickangle=-45)
            st.plotly_chart(fig, use_container_width=True)
        else:
            st.info("Este municipio no tiene proyectos con anexos ISC aplicables.")
            
        st.subheader("Detalle de proyectos del municipio")
        cols_resumen = [
            "codigo_interno", "municipio", "codigo_sigesplan", "linea_inversion",
            "tipo_proyecto", "titulo_proyecto", "profesional_spae",
            "profesional_ggp", "presupuesto_proyectado",
            "presupuesto_solicitado", "total_beneficiados",
            "pct_anexos_isc", "categoria_isc", "estado_formulacion",
        ]
        cols_existentes = [c for c in cols_resumen if c in df_muni.columns]
        st.dataframe(df_muni[cols_existentes], use_container_width=True, hide_index=True)

with tab_grupo:
    st.markdown("## Análisis de Portafolio por Equipo / Grupo")
    grupo_seleccionado = st.radio("Seleccione el grupo a analizar:", options=["SPAE", "GGP"], horizontal=True)
    col_responsable = "profesional_spae" if grupo_seleccionado == "SPAE" else "profesional_ggp"
    
    if col_responsable in df.columns:
        df_valid = df.dropna(subset=[col_responsable])
        df_valid = df_valid[df_valid[col_responsable].astype(str).str.strip() != ""]
        
        if not df_valid.empty:
            df_valid["presupuesto_proyectado"] = pd.to_numeric(df_valid["presupuesto_proyectado"], errors='coerce').fillna(0)
            df_valid["total_beneficiados"] = pd.to_numeric(df_valid["total_beneficiados"], errors='coerce').fillna(0)
            df_valid["pct_anexos_isc"] = pd.to_numeric(df_valid.get("pct_anexos_isc", 0), errors='coerce').fillna(0)
            
            df_group = df_valid.groupby(col_responsable).agg(
                Proyectos_Asignados=("codigo_proyecto", "count"),
                Presupuesto_A_Cargo=("presupuesto_proyectado", "sum"),
                Beneficiados_A_Cargo=("total_beneficiados", "sum"),
                Promedio_Anexos_Recibidos=("pct_anexos_isc", "mean")
            ).reset_index()
            
            df_group["Promedio_Anexos_Recibidos"] = df_group["Promedio_Anexos_Recibidos"].round(1)
            
            st.markdown(f"### Resumen Global Equipo {grupo_seleccionado}")
            t1, t2, t3, t4 = st.columns(4)
            t1.metric(f"Total Profesionales", len(df_group))
            t2.metric(f"Total Proyectos Asignados", df_group["Proyectos_Asignados"].sum())
            t3.metric(f"Presupuesto Bajo Responsabilidad", f"${df_group['Presupuesto_A_Cargo'].sum():,.0f}")
            t4.metric(f"Promedio Avance Anexos", f"{df_group['Promedio_Anexos_Recibidos'].mean():.1f}%")
            
            st.divider()
            
            st.markdown("### Cargas y Rendimiento por Profesional")
            st.dataframe(
                df_group.sort_values("Proyectos_Asignados", ascending=False), 
                use_container_width=True, hide_index=True,
                column_config={
                    "Presupuesto_A_Cargo": st.column_config.NumberColumn("Presupuesto a Cargo", format="$%d"),
                    "Promedio_Anexos_Recibidos": st.column_config.ProgressColumn("Avance % Anexos ISC", format="%f", min_value=0, max_value=100)
                }
            )
            
            fig_carga = px.bar(df_group, x=col_responsable, y="Proyectos_Asignados", title=f"Proyectos Asignados por Profesional ({grupo_seleccionado})", color="Promedio_Anexos_Recibidos", color_continuous_scale="Viridis")
            st.plotly_chart(fig_carga, use_container_width=True)
        else:
            st.info(f"No hay profesionales registrados bajo el grupo {grupo_seleccionado}.")

with tab_crm:
    st.markdown("## Directorio Global e Historial CRM")
    
    with st.expander(" Subir Directorio de Profesionales a Supabase"):
        st.info("Carga el archivo Excel con el Directorio de Profesionales Humanitarios (00. Directorio...).")
        uploaded_directorio = st.file_uploader("Arrastra el Directorio (.xlsx)", type=["xlsx"], key="upload_dir_prof")
        if uploaded_directorio:
            if st.button("Subir a la nube"):
                import pandas as pd
                import json
                from auth import save_project_to_db
                try:
                    df_dir_prof = pd.read_excel(uploaded_directorio, header=None).fillna("")
                    estado_str = json.dumps(df_dir_prof.to_dict(orient='records'), ensure_ascii=False)
                    user = st.session_state.get("auth_user")
                    if not user:
                        st.error("No hay un usuario autenticado. Inicia sesión nuevamente.")
                    else:
                        save_project_to_db(user, "Directorio Profesionales Humanitarios - 2026", "admin", "", "", estado_str)
                        st.success("¡Directorio subido exitosamente a Supabase!")
                except Exception as e:
                    st.error(f"Error al subir: {e}")
    with st.expander("Ver Directorio de Profesionales Humanitarios (Nube)"):
        try:
            user = st.session_state.get("auth_user")
            if user:
                from auth import SUPABASE_URL, SUPABASE_KEY
                import requests
                import json
                import pandas as pd
                
                url = f"{SUPABASE_URL}/rest/v1/proyectos?nombre_proyecto=eq.Directorio Profesionales Humanitarios - 2026&select=estado_json&order=created_at.desc&limit=1"
                headers = {"apikey": SUPABASE_KEY, "Authorization": f"Bearer {getattr(user, 'access_token', SUPABASE_KEY)}"}
                res = requests.get(url, headers=headers)
                if res.status_code == 200:
                    data = res.json()
                    if data and len(data) > 0:
                        estado = data[0].get("estado_json", "[]")
                        if isinstance(estado, str): estado = json.loads(estado)
                        df_dir_prof = pd.DataFrame(estado)
                        
                        # Asignar nombres legibles a las columnas (originalmente eran 0, 1, 2, 3, 4, 5)
                        if len(df_dir_prof.columns) == 6:
                            df_dir_prof.columns = ["Territorial", "Nombre", "Correo", "Teléfono", "Cargo", "Dirección"]
                        
                        import urllib.parse
                        import re
                        
                        def gen_wa_prof(tel):
                            tel_str = str(tel).strip()
                            if not tel_str or tel_str in ["", "nan", "None"]: return None
                            digits = re.sub(r'\D', '', tel_str)
                            if not digits: return None
                            if not digits.startswith("57"): digits = "57" + digits
                            return f"https://wa.me/{digits}"
                            
                        def gen_email_prof(correo, nombre):
                            c_str = str(correo).strip()
                            if "@" not in c_str: return None
                            asunto = "Contacto desde Portafolio SPAE - Directorio de Profesionales"
                            cuerpo = (
                                f"Cordial saludo {nombre},\n\n"
                                "Me comunico con usted para tratar temas de articulación y seguimiento "
                                "relacionados con los proyectos SPAE de su departamento.\n\n"
                                "[Escriba su mensaje aquí...]\n\n"
                                "Quedo atento(a).\n\n"
                                "Profesional SPAE"
                            )
                            return f"https://mail.google.com/mail/?view=cm&fs=1&to={c_str}&su={urllib.parse.quote(asunto)}&body={urllib.parse.quote(cuerpo)}"

                        if "Teléfono" in df_dir_prof.columns:
                            df_dir_prof["Contactar (WhatsApp)"] = df_dir_prof["Teléfono"].apply(gen_wa_prof)
                        if "Correo" in df_dir_prof.columns and "Nombre" in df_dir_prof.columns:
                            df_dir_prof["Enviar Email (Gmail)"] = df_dir_prof.apply(lambda r: gen_email_prof(r.get("Correo"), r.get("Nombre", "")), axis=1)
                        
                        col_cfg_prof = {
                            "Contactar (WhatsApp)": st.column_config.LinkColumn("Contactar (WhatsApp)", display_text="Abrir Chat WA"),
                            "Enviar Email (Gmail)": st.column_config.LinkColumn("Enviar Email (Gmail)", display_text="Abrir Gmail")
                        }
                        
                        # Filtrar por los departamentos asignados al usuario / municipio seleccionado
                        import unicodedata
                        def normalize_text(text):
                            if not isinstance(text, str): return ""
                            return unicodedata.normalize('NFD', text).encode('ascii', 'ignore').decode('utf-8').lower().strip()
                        
                        try:
                            if not df_muni.empty and "departamento" in df_muni.columns:
                                deptos_muni = [normalize_text(d) for d in df_muni["departamento"].dropna().unique()]
                                if deptos_muni:
                                    def match_depto(territorial_val):
                                        t_norm = normalize_text(territorial_val)
                                        for d in deptos_muni:
                                            if d in t_norm: return True
                                        return False
                                    df_dir_prof = df_dir_prof[df_dir_prof["Territorial"].apply(match_depto)]
                                    if df_dir_prof.empty:
                                        st.warning("No se encontraron profesionales en el directorio para tus departamentos asignados.")
                        except Exception as e:
                            pass # Fallback a mostrar todos si hay error en el filtrado
                        
                        st.dataframe(df_dir_prof, use_container_width=True, hide_index=True, column_config=col_cfg_prof)
                    else:
                        st.info("Aún no has subido el Directorio de Profesionales a la nube.")
                else:
                    st.error("Error consultando la base de datos.")
            else:
                st.warning("Debes iniciar sesión para ver el directorio.")
        except Exception as e:
            st.error(f"Error al cargar: {e}")
            
    if not df_muni.empty and "municipio" in df_muni.columns:
        df_display = df_muni.copy()
        
        wa_links, wa_sources, info_contactos = [], [], []
        for _, row in df_display.iterrows():
            link, source = whatsapp_link_cascada(row)
            wa_links.append(link)
            wa_sources.append(source)
            tel = row.get("telefono_enlace") if pd.notna(row.get("telefono_enlace")) else row.get("telefono_alcalde")
            info_contactos.append(str(tel) if pd.notna(tel) and str(tel).strip() not in ["", "0", "None", "nan"] else "Sin Información")
            
        df_display["Contactar (WhatsApp)"] = wa_links
        df_display["Fuente del Contacto"] = wa_sources
        df_display["Info Contacto"] = info_contactos
        
        if "correos_alcaldia" in df_display.columns:
            df_display["Correo Contacto"] = df_display["correos_alcaldia"].apply(
                lambda x: str(x).split()[0].split(',')[0].strip() if pd.notna(x) and "@" in str(x) else "Sin Correo"
            )
            df_display["Enviar Email (Auto)"] = df_display.apply(lambda r: email_link(r), axis=1)
        else:
            df_display["Correo Contacto"] = "Sin Correo"
            df_display["Enviar Email (Auto)"] = None

        with st.expander("Agregar / Actualizar Contacto en Directorio Maestro", expanded=False):
            col1, col2 = st.columns(2)
            with col1:
                nuevo_muni = st.text_input("Municipio del contacto (ej. Sipí)")
                nuevo_enlace = st.text_input("Nombre y Teléfono del Enlace (ej. Juan Perez 3001234567)")
            with col2:
                nuevo_alcalde = st.text_input("Nombre y Teléfono del Alcalde (Opcional)")
                nuevo_correo = st.text_input("Correo de Alcaldía (Opcional)")
            
            if st.button("Guardar Contacto en Excel", use_container_width=True):
                if nuevo_muni.strip():
                    ruta_excel = r"C:\Users\cagch\Desktop\senador\upv2026\junio\Matriz_Seguimiento_Proyectos_Cesar.xlsx"
                    try:
                        import pandas as pd
                        df_update = pd.read_excel(ruta_excel)
                        muni_norm = nuevo_muni.strip().lower()
                        idx_match = None
                        for idx, r_ in df_update.iterrows():
                            if str(r_.get("Municipio")).strip().lower() == muni_norm:
                                idx_match = idx
                                break
                        if idx_match is not None:
                            if nuevo_enlace: df_update.at[idx_match, "Teléfono Enlace"] = nuevo_enlace
                            if nuevo_alcalde: df_update.at[idx_match, "Teléfono Alcaldía"] = nuevo_alcalde
                            if nuevo_correo: df_update.at[idx_match, "Correos Alcaldía"] = nuevo_correo
                            st.success(f"Contacto actualizado para {nuevo_muni}. Recargando...")
                        else:
                            nueva_fila = {"Municipio": nuevo_muni.strip()}
                            if nuevo_enlace: nueva_fila["Teléfono Enlace"] = nuevo_enlace
                            if nuevo_alcalde: nueva_fila["Teléfono Alcaldía"] = nuevo_alcalde
                            if nuevo_correo: nueva_fila["Correos Alcaldía"] = nuevo_correo
                            df_update = pd.concat([df_update, pd.DataFrame([nueva_fila])], ignore_index=True)
                            st.success(f"Nuevo municipio {nuevo_muni} agregado al directorio. Recargando...")
                        df_update.to_excel(ruta_excel, index=False)
                        st.rerun()
                    except Exception as e:
                        st.error(f"Error guardando contacto: {e}")
                else:
                    st.warning("El nombre del municipio es obligatorio.")
        
        cols_to_show = ["municipio", "nombre_proyecto", "estado_formulacion", "profesional_spae", "profesional_ggp", 
                        "Info Contacto", "Fuente del Contacto", "Contactar (WhatsApp)", "Correo Contacto", "Enviar Email (Auto)"]
        
        col_config = {
            "Contactar (WhatsApp)": st.column_config.LinkColumn("Contactar (WhatsApp)", display_text="Abrir Chat WA"),
            "Enviar Email (Auto)": st.column_config.LinkColumn("Enviar Email (Auto)", display_text="Abrir App Correo")
        }

        cols_final = [c for c in cols_to_show if c in df_display.columns]

        st.dataframe(
            df_display[cols_final],
            use_container_width=True,
            hide_index=True,
            column_config=col_config
        )

with tab_historial:
    st.markdown("## Historial de Revisiones Local (JSON)")
    df_hist = get_local_history()
    if not df_hist.empty:
        st.dataframe(df_hist, hide_index=True, use_container_width=True)
    else:
        st.info("No hay historial local previo. A partir de hoy se registrarán aquí los cambios que subas.")


# ------------------------------------------------------------------------------
# PANEL DE GENERACION Y DESCARGA DE INFORMES EN WORD
# ------------------------------------------------------------------------------
st.sidebar.markdown("---")
st.sidebar.subheader("Modulo de Informes y Descargas")

if st.sidebar.button("Generar Informe General del Portafolio"):
    try:
        df_gen = pd.read_excel(r"C:\Users\cagch\Desktop\SPAE_Aislada\data\Matriz_Seguimiento_Proyectos_Cesar.xlsx")
        docx_bytes = generar_documento_word_informe(df_gen, "Informe General de Seguimiento Portafolio SPAE 2026")
        st.sidebar.download_button(
            label="Descargar Informe General (Word)",
            data=docx_bytes,
            file_name="Informe_General_Portafolio_SPAE_2026.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
    except Exception as e_err:
        st.sidebar.error(f"Error generando informe: {e_err}")

if st.sidebar.button("Generar Informe Mis Proyectos (Cesar Giraldo)"):
    try:
        df_all = pd.read_excel(r"C:\Users\cagch\Desktop\SPAE_Aislada\data\Matriz_Seguimiento_Proyectos_Cesar.xlsx")
        df_cesar = df_all[df_all["Profesional SPAE"].astype(str).str.contains("Cesar", case=False, na=False)]
        docx_bytes_c = generar_documento_word_informe(df_cesar, "Informe de Gestion - Ing. Cesar Giraldo Chaparro (11 Municipios)")
        st.sidebar.download_button(
            label="Descargar Informe Mis Proyectos (Word)",
            data=docx_bytes_c,
            file_name="Informe_Mis_Proyectos_Cesar_Giraldo_11_Municipios.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
    except Exception as e_err:
        st.sidebar.error(f"Error generando informe: {e_err}")

st.sidebar.markdown("---")
